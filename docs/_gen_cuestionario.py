# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import qn

path = r"c:\Users\kyrian\Documents\06-Scout\Scout\Encuentro_Rover\docs\Cuestionario-Emilio-Aguero.docx"
doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.2)
    s.right_margin = Cm(2.2)


def font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def para(text, *, bold=False, italic=False, size=11, color=None, after=6, before=0):
    p = doc.add_paragraph()
    r = p.add_run(text)
    font(r, size, bold, color)
    r.italic = italic
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)


def category(title):
    p = doc.add_paragraph()
    r = p.add_run(title)
    font(r, 14, True, RGBColor(0x1E, 0x5A, 0x8A))
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)


def question(num, title, details):
    p = doc.add_paragraph()
    r = p.add_run(f"{num}. {title}")
    font(r, 12, True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)

    if details:
        para(details, italic=True, size=10, color=RGBColor(0x5C, 0x6B, 0x7A), after=6)

    for _ in range(4):
        blank = doc.add_paragraph()
        blank.paragraph_format.space_after = Pt(8)


p = doc.add_paragraph()
r = p.add_run("Cuestionario para Emilio Agüero")
font(r, 18, True, RGBColor(0x1A, 0x36, 0x5D))
p.paragraph_format.space_after = Pt(6)

para("Justa del Saber — Encuentro Rover", bold=True, size=12)
para(
    "Respondé debajo de cada pregunta. Cuando termines, guardá y devolvé este archivo por WhatsApp.",
    size=11,
    after=8,
)
para("Nombre: Emilio Agüero", size=11)
para("Fecha: _______________", size=11, after=4)

# Grouped related questions — fewer items, related asks together in the detail line
items = [
    (
        "1. Clanes",
        [
            (
                "Clanes participantes",
                "¿Qué clanes van a participar, cuántos son, y tenés la lista en Excel/CSV? ¿Cuándo la podés mandar?",
            ),
        ],
    ),
    (
        "2. Logos",
        [
            (
                "Logos de los clanes",
                "¿Existen logos? ¿En qué formato? Si faltan, ¿está bien usar iniciales y un color por clan?",
            ),
        ],
    ),
    (
        "3. Preguntas del juego",
        [
            (
                "Banco de preguntas",
                "¿Cuántas preguntas va a tener el juego? ¿Ya están escritas? ¿Incluyen la respuesta correcta para el host? ¿Quién las arma/aprueba y para cuándo las podés mandar?",
            ),
        ],
    ),
    (
        "4. Participantes",
        [
            (
                "Lista de participantes",
                "¿Tenés la lista de asistencia? ¿Hace falta usarla el día del evento o solo es registro? (En el juego compiten los clanes.)",
            ),
        ],
    ),
    (
        "5. Título y pantalla",
        [
            (
                "Título en pantalla",
                "¿Qué título debe ver el público? (ej. Justa del Saber, Encuentro Rover, u otro)",
            ),
            (
                "Respuesta correcta del host",
                "El host usa la misma pantalla proyectada. ¿Cómo ve la respuesta oficial sin que el público la lea?",
            ),
        ],
    ),
    (
        "6. Reglas del juego",
        [
            (
                "Flujo general",
                "Propuesta: ruleta elige clan + pregunta → timer 60 s → responden en voz alta → host marca correcto/incorrecto (o prórroga; si vence el tiempo → 0) → se revela la respuesta → puntajes → al final tabla. Preguntas no se repiten; clanes sí pueden repetir. El host puede cerrar antes. ¿Está bien o qué cambiarías?",
            ),
            (
                "Puntaje",
                "Propuesta: Correcto = 10 + segundos que quedan; Incorrecto = 0; timer 60 s. ¿OK o qué números querés?",
            ),
            (
                "Empates",
                "Si dos clanes terminan con los mismos puntos, ¿qué pasa?",
            ),
        ],
    ),
    (
        "7. Arma secreta",
        [
            (
                "Arma secreta",
                "En tu audio con Dani hablaste de un “arma secreta”. ¿Es esta Justa del Saber, o hay algo más?",
            ),
        ],
    ),
    (
        "8. Logística del día",
        [
            (
                "Equipo e internet",
                "¿Hay proyector y notebook estables? ¿Quién se encarga? ¿Habrá internet (GitHub Pages) o preferís que funcione offline?",
            ),
            (
                "Host y ensayo",
                "¿Quién va a ser el host/operador en vivo? ¿Hay ensayo previo? ¿Cuándo?",
            ),
        ],
    ),
    (
        "9. Algo más",
        [
            (
                "Otros detalles",
                "¿Hay alguna otra regla o detalle importante que no hayamos preguntado?",
            ),
        ],
    ),
]

n = 0
for cat, qs in items:
    category(cat)
    for title, details in qs:
        n += 1
        question(n, title, details)

para("¡Gracias! Guardá y devolvé este archivo.", bold=True, size=12, before=16)

doc.save(path)
print("Saved:", path)
print("Questions:", n)
print("Categories:", len(items))
